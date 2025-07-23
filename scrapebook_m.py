# Word Document Exploratory Data Analysis (EDA) with Python
# A comprehensive notebook for analyzing Word documents

# ## 1. Setup and Installation

# First, install required packages (run this cell if packages are not installed)
# ```bash
# pip install python-docx pandas matplotlib seaborn wordcloud nltk textblob textstat scikit-learn plotly
# pip install textract  # For handling .doc files and other formats
# ```

# ## 2. Import Required Libraries

import docx
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
from collections import Counter
import re
from pathlib import Path
import warnings
warnings.filterwarnings('ignore')

# Text processing libraries
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.stem import WordNetLemmatizer
from textblob import TextBlob
import textstat

# Visualization libraries
from wordcloud import WordCloud
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots

# Machine learning libraries
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.decomposition import LatentDirichletAllocation
from sklearn.cluster import KMeans

# Download NLTK data (run once)
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
    nltk.data.find('corpora/wordnet')
except LookupError:
    nltk.download('punkt')
    nltk.download('stopwords')
    nltk.download('wordnet')
    nltk.download('averaged_perceptron_tagger')

print("All libraries imported successfully!")

# ## 3. Document Loading and Text Extraction Functions

class WordDocumentAnalyzer:
    """A comprehensive class for analyzing Word documents"""
    
    def __init__(self):
        self.documents = {}
        self.processed_docs = {}
        self.doc_stats = pd.DataFrame()
        
    def load_single_document(self, file_path):
        """Extract text from a single .docx file"""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    tables_text.append(' | '.join(row_text))
            
            # Combine all text
            full_text = '\n'.join(paragraphs)
            if tables_text:
                full_text += '\n\nTables:\n' + '\n'.join(tables_text)
                
            return full_text
            
        except Exception as e:
            print(f"Error loading {file_path}: {e}")
            return ""
    
    def load_multiple_documents(self, file_paths):
        """Load multiple Word documents"""
        self.documents = {}
        
        for file_path in file_paths:
            doc_name = Path(file_path).stem
            text = self.load_single_document(file_path)
            if text:
                self.documents[doc_name] = text
                print(f"✓ Loaded: {doc_name}")
            else:
                print(f"✗ Failed to load: {doc_name}")
        
        print(f"\nTotal documents loaded: {len(self.documents)}")
        return self.documents
    
    def load_from_directory(self, directory_path, file_extension="*.docx"):
        """Load all Word documents from a directory"""
        directory = Path(directory_path)
        file_paths = list(directory.glob(file_extension))
        
        if not file_paths:
            print(f"No {file_extension} files found in {directory_path}")
            return {}
            
        return self.load_multiple_documents(file_paths)

# Initialize the analyzer
analyzer = WordDocumentAnalyzer()

# ## 4. Sample Data Creation (for demonstration)

# Let's create some sample documents for demonstration
sample_docs = {
    "sample_doc_1": """
    This is a sample business report about quarterly performance. 
    Our company has shown significant growth this quarter with revenue increasing by 25%.
    The marketing team has done excellent work in expanding our customer base.
    We need to focus more on customer satisfaction and product quality.
    The next quarter looks promising with several new product launches planned.
    """,
    
    "sample_doc_2": """
    Technical documentation for the new software system.
    The system architecture follows microservices pattern with Docker containers.
    Database optimization has improved query performance by 40%.
    Security measures include authentication, authorization, and data encryption.
    Testing coverage has reached 95% with automated test suites.
    Future developments will focus on scalability and performance enhancements.
    """,
    
    "sample_doc_3": """
    Research findings on market trends and consumer behavior.
    Consumer preferences have shifted towards sustainable products.
    Digital transformation has accelerated across all industries.
    Remote work has become the new normal affecting workplace dynamics.
    E-commerce growth continues to outpace traditional retail channels.
    Data privacy concerns are influencing purchasing decisions.
    """
}

# Use sample documents for demonstration
analyzer.documents = sample_docs
print("Sample documents loaded for demonstration.")

# ## 5. Basic Document Statistics

def calculate_document_statistics(text):
    """Calculate comprehensive statistics for a document"""
    
    # Basic counts
    words = text.split()
    sentences = sent_tokenize(text)
    paragraphs = [p for p in text.split('\n') if p.strip()]
    
    # Character analysis
    chars_total = len(text)
    chars_no_spaces = len(text.replace(' ', ''))
    
    # Word analysis
    word_lengths = [len(word) for word in words]
    unique_words = set(word.lower().strip('.,!?";') for word in words)
    
    # Sentence analysis
    sentence_lengths = [len(sent.split()) for sent in sentences]
    
    stats = {
        'total_characters': chars_total,
        'characters_no_spaces': chars_no_spaces,
        'total_words': len(words),
        'unique_words': len(unique_words),
        'total_sentences': len(sentences),
        'total_paragraphs': len(paragraphs),
        'avg_word_length': np.mean(word_lengths) if word_lengths else 0,
        'avg_sentence_length': np.mean(sentence_lengths) if sentence_lengths else 0,
        'avg_words_per_paragraph': len(words) / len(paragraphs) if paragraphs else 0,
        'lexical_diversity': len(unique_words) / len(words) if words else 0
    }
    
    return stats

# Calculate statistics for all documents
stats_data = {}
for doc_name, text in analyzer.documents.items():
    stats_data[doc_name] = calculate_document_statistics(text)

# Create DataFrame
analyzer.doc_stats = pd.DataFrame(stats_data).T
print("Document Statistics:")
print(analyzer.doc_stats.round(2))

# ## 6. Text Preprocessing Functions

class TextPreprocessor:
    """Class for text preprocessing operations"""
    
    def __init__(self):
        self.stop_words = set(stopwords.words('english'))
        self.lemmatizer = WordNetLemmatizer()
        
    def clean_text(self, text):
        """Basic text cleaning"""
        # Convert to lowercase
        text = text.lower()
        
        # Remove special characters but keep spaces and periods
        text = re.sub(r'[^\w\s.]', ' ', text)
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    def tokenize_and_preprocess(self, text):
        """Tokenize and preprocess text"""
        # Clean text first
        cleaned_text = self.clean_text(text)
        
        # Tokenize
        tokens = word_tokenize(cleaned_text)
        
        # Remove stopwords and short words
        tokens = [token for token in tokens if token not in self.stop_words and len(token) > 2]
        
        # Lemmatize
        tokens = [self.lemmatizer.lemmatize(token) for token in tokens]
        
        return tokens
    
    def get_word_frequency(self, tokens):
        """Get word frequency distribution"""
        return Counter(tokens)

# Initialize preprocessor
preprocessor = TextPreprocessor()

# Process all documents
processed_docs = {}
word_frequencies = {}

for doc_name, text in analyzer.documents.items():
    tokens = preprocessor.tokenize_and_preprocess(text)
    processed_docs[doc_name] = tokens
    word_frequencies[doc_name] = preprocessor.get_word_frequency(tokens)

analyzer.processed_docs = processed_docs
print("Text preprocessing completed.")

# ## 7. Visualization Functions

def create_document_stats_dashboard():
    """Create an interactive dashboard for document statistics"""
    
    # Create subplots
    fig = make_subplots(
        rows=2, cols=2,
        subplot_titles=('Word Count', 'Sentence Count', 'Average Word Length', 'Lexical Diversity'),
        specs=[[{"secondary_y": False}, {"secondary_y": False}],
               [{"secondary_y": False}, {"secondary_y": False}]]
    )
    
    doc_names = analyzer.doc_stats.index.tolist()
    
    # Word count
    fig.add_trace(
        go.Bar(x=doc_names, y=analyzer.doc_stats['total_words'], name='Word Count', marker_color='lightblue'),
        row=1, col=1
    )
    
    # Sentence count
    fig.add_trace(
        go.Bar(x=doc_names, y=analyzer.doc_stats['total_sentences'], name='Sentence Count', marker_color='lightgreen'),
        row=1, col=2
    )
    
    # Average word length
    fig.add_trace(
        go.Bar(x=doc_names, y=analyzer.doc_stats['avg_word_length'], name='Avg Word Length', marker_color='lightsalmon'),
        row=2, col=1
    )
    
    # Lexical diversity
    fig.add_trace(
        go.Bar(x=doc_names, y=analyzer.doc_stats['lexical_diversity'], name='Lexical Diversity', marker_color='lightpink'),
        row=2, col=2
    )
    
    fig.update_layout(height=600, showlegend=False, title_text="Document Statistics Dashboard")
    fig.show()

# Create the dashboard
create_document_stats_dashboard()

# ## 8. Word Frequency Analysis

def plot_word_frequency_comparison(top_n=10):
    """Plot word frequency comparison across documents"""
    
    fig, axes = plt.subplots(1, len(analyzer.documents), figsize=(15, 5))
    if len(analyzer.documents) == 1:
        axes = [axes]
    
    for idx, (doc_name, freq) in enumerate(word_frequencies.items()):
        top_words = dict(freq.most_common(top_n))
        
        axes[idx].bar(range(len(top_words)), list(top_words.values()), color=f'C{idx}')
        axes[idx].set_xticks(range(len(top_words)))
        axes[idx].set_xticklabels(list(top_words.keys()), rotation=45, ha='right')
        axes[idx].set_title(f'Top {top_n} Words: {doc_name}')
        axes[idx].set_ylabel('Frequency')
    
    plt.tight_layout()
    plt.show()

plot_word_frequency_comparison(10)

# ## 9. Word Cloud Generation

def generate_word_clouds():
    """Generate word clouds for all documents"""
    
    n_docs = len(analyzer.documents)
    cols = min(3, n_docs)
    rows = (n_docs + cols - 1) // cols
    
    fig, axes = plt.subplots(rows, cols, figsize=(15, 5*rows))
    if n_docs == 1:
        axes = [axes]
    elif rows == 1:
        axes = axes.reshape(1, -1)
    
    # Flatten axes for easier iteration
    axes_flat = axes.flatten() if n_docs > 1 else axes
    
    for idx, (doc_name, text) in enumerate(analyzer.documents.items()):
        if idx < len(axes_flat):
            wordcloud = WordCloud(
                width=400, height=300,
                background_color='white',
                max_words=100,
                colormap='viridis'
            ).generate(text)
            
            axes_flat[idx].imshow(wordcloud, interpolation='bilinear')
            axes_flat[idx].set_title(f'Word Cloud: {doc_name}')
            axes_flat[idx].axis('off')
    
    # Hide empty subplots
    for idx in range(n_docs, len(axes_flat)):
        axes_flat[idx].axis('off')
    
    plt.tight_layout()
    plt.show()

generate_word_clouds()

# ## 10. Sentiment Analysis

def perform_sentiment_analysis():
    """Perform sentiment analysis on all documents"""
    
    sentiment_data = {}
    
    for doc_name, text in analyzer.documents.items():
        blob = TextBlob(text)
        
        sentiment_data[doc_name] = {
            'polarity': blob.sentiment.polarity,
            'subjectivity': blob.sentiment.subjectivity,
            'polarity_label': 'Positive' if blob.sentiment.polarity > 0 else 'Negative' if blob.sentiment.polarity < 0 else 'Neutral'
        }
    
    sentiment_df = pd.DataFrame(sentiment_data).T
    
    # Create visualization
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 5))
    
    # Polarity plot
    colors = ['red' if p < 0 else 'green' if p > 0 else 'gray' for p in sentiment_df['polarity']]
    ax1.bar(sentiment_df.index, sentiment_df['polarity'], color=colors, alpha=0.7)
    ax1.set_title('Sentiment Polarity by Document')
    ax1.set_ylabel('Polarity (-1 to 1)')
    ax1.axhline(y=0, color='black', linestyle='-', alpha=0.3)
    ax1.tick_params(axis='x', rotation=45)
    
    # Subjectivity plot
    ax2.bar(sentiment_df.index, sentiment_df['subjectivity'], color='blue', alpha=0.7)
    ax2.set_title('Subjectivity by Document')
    ax2.set_ylabel('Subjectivity (0 to 1)')
    ax2.tick_params(axis='x', rotation=45)
    
    plt.tight_layout()
    plt.show()
    
    print("Sentiment Analysis Results:")
    print(sentiment_df.round(3))
    
    return sentiment_df

sentiment_results = perform_sentiment_analysis()

# ## 11. Readability Analysis

def analyze_readability():
    """Analyze readability metrics for all documents"""
    
    readability_data = {}
    
    for doc_name, text in analyzer.documents.items():
        readability_data[doc_name] = {
            'flesch_kincaid_grade': textstat.flesch_kincaid_grade(text),
            'flesch_reading_ease': textstat.flesch_reading_ease(text),
            'automated_readability_index': textstat.automated_readability_index(text),
            'coleman_liau_index': textstat.coleman_liau_index(text),
            'gunning_fog': textstat.gunning_fog(text)
        }
    
    readability_df = pd.DataFrame(readability_data).T
    
    # Create visualization
    fig, axes = plt.subplots(2, 3, figsize=(18, 10))
    axes = axes.flatten()
    
    metrics = readability_df.columns
    
    for idx, metric in enumerate(metrics):
        axes[idx].bar(readability_df.index, readability_df[metric], alpha=0.7, color=f'C{idx}')
        axes[idx].set_title(f'{metric.replace("_", " ").title()}')
        axes[idx].tick_params(axis='x', rotation=45)
        
    # Hide the last empty subplot
    axes[-1].axis('off')
    
    plt.tight_layout()
    plt.show()
    
    print("Readability Analysis Results:")
    print(readability_df.round(2))
    
    return readability_df

readability_results = analyze_readability()

# ## 12. Topic Modeling with LDA

def perform_topic_modeling(n_topics=3, n_words=10):
    """Perform topic modeling using Latent Dirichlet Allocation"""
    
    # Prepare documents for topic modeling
    documents_text = list(analyzer.documents.values())
    
    if len(documents_text) < 2:
        print("Need at least 2 documents for meaningful topic modeling.")
        return None
    
    # Vectorize the documents
    vectorizer = TfidfVectorizer(
        max_features=100,
        stop_words='english',
        lowercase=True,
        token_pattern=r'\b[a-zA-Z]{3,}\b'
    )
    
    doc_term_matrix = vectorizer.fit_transform(documents_text)
    feature_names = vectorizer.get_feature_names_out()
    
    # Perform LDA
    lda = LatentDirichletAllocation(
        n_components=n_topics,
        random_state=42,
        max_iter=10
    )
    lda.fit(doc_term_matrix)
    
    # Extract topics
    topics = []
    for topic_idx, topic in enumerate(lda.components_):
        top_words_idx = topic.argsort()[-n_words:][::-1]
        top_words = [feature_names[i] for i in top_words_idx]
        topics.append(top_words)
    
    # Create visualization
    fig, ax = plt.subplots(figsize=(12, 6))
    
    y_pos = np.arange(n_topics)
    topic_labels = [f"Topic {i+1}: {', '.join(topic[:5])}" for i, topic in enumerate(topics)]
    
    # Calculate topic strengths (sum of word weights)
    topic_strengths = [np.sum(lda.components_[i]) for i in range(n_topics)]
    
    bars = ax.barh(y_pos, topic_strengths, alpha=0.7)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(topic_labels)
    ax.set_xlabel('Topic Strength')
    ax.set_title('Discovered Topics')
    
    # Add value labels on bars
    for i, bar in enumerate(bars):
        width = bar.get_width()
        ax.text(width + 0.01, bar.get_y() + bar.get_height()/2, 
                f'{width:.2f}', ha='left', va='center')
    
    plt.tight_layout()
    plt.show()
    
    # Print detailed topics
    print("Discovered Topics:")
    for i, topic_words in enumerate(topics):
        print(f"Topic {i+1}: {', '.join(topic_words)}")
    
    return topics, lda, vectorizer

topics, lda_model, vectorizer = perform_topic_modeling()

# ## 13. Document Similarity Analysis

def analyze_document_similarity():
    """Analyze similarity between documents using TF-IDF"""
    
    if len(analyzer.documents) < 2:
        print("Need at least 2 documents for similarity analysis.")
        return None
    
    # Create TF-IDF vectors
    documents_text = list(analyzer.documents.values())
    doc_names = list(analyzer.documents.keys())
    
    tfidf_vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)
    tfidf_matrix = tfidf_vectorizer.fit_transform(documents_text)
    
    # Calculate cosine similarity
    from sklearn.metrics.pairwise import cosine_similarity
    
    similarity_matrix = cosine_similarity(tfidf_matrix)
    
    # Create similarity DataFrame
    similarity_df = pd.DataFrame(
        similarity_matrix,
        index=doc_names,
        columns=doc_names
    )
    
    # Create heatmap
    plt.figure(figsize=(10, 8))
    sns.heatmap(
        similarity_df,
        annot=True,
        cmap='coolwarm',
        center=0,
        square=True,
        fmt='.3f'
    )
    plt.title('Document Similarity Matrix')
    plt.tight_layout()
    plt.show()
    
    print("Document Similarity Matrix:")
    print(similarity_df.round(3))
    
    return similarity_df

similarity_matrix = analyze_document_similarity()

# ## 14. Comprehensive Report Generation

def generate_comprehensive_report():
    """Generate a comprehensive EDA report"""
    
    print("="*60)
    print("COMPREHENSIVE WORD DOCUMENT EDA REPORT")
    print("="*60)
    
    print(f"\n📊 DATASET OVERVIEW")
    print(f"Total Documents: {len(analyzer.documents)}")
    print(f"Document Names: {list(analyzer.documents.keys())}")
    
    print(f"\n📈 BASIC STATISTICS")
    print(analyzer.doc_stats.round(2))
    
    if len(analyzer.documents) > 1:
        print(f"\n🔍 DOCUMENT COMPARISON")
        print(f"Most words: {analyzer.doc_stats['total_words'].idxmax()} ({analyzer.doc_stats['total_words'].max()} words)")
        print(f"Least words: {analyzer.doc_stats['total_words'].idxmin()} ({analyzer.doc_stats['total_words'].min()} words)")
        print(f"Highest lexical diversity: {analyzer.doc_stats['lexical_diversity'].idxmax()} ({analyzer.doc_stats['lexical_diversity'].max():.3f})")
        print(f"Average document length: {analyzer.doc_stats['total_words'].mean():.1f} words")
    
    print(f"\n💭 SENTIMENT ANALYSIS")
    if 'sentiment_results' in globals():
        for doc_name, sentiment in sentiment_results.iterrows():
            print(f"{doc_name}: {sentiment['polarity_label']} (Polarity: {sentiment['polarity']:.3f}, Subjectivity: {sentiment['subjectivity']:.3f})")
    
    print(f"\n📚 READABILITY ANALYSIS")
    if 'readability_results' in globals():
        avg_grade_level = readability_results['flesch_kincaid_grade'].mean()
        print(f"Average Grade Level: {avg_grade_level:.1f}")
        print(f"Average Reading Ease: {readability_results['flesch_reading_ease'].mean():.1f}")
    
    print(f"\n🏷️ TOP WORDS ACROSS ALL DOCUMENTS")
    # Combine all word frequencies
    all_words = Counter()
    for freq in word_frequencies.values():
        all_words.update(freq)
    
    top_words = all_words.most_common(10)
    for word, count in top_words:
        print(f"{word}: {count}")
    
    if len(analyzer.documents) > 1 and 'topics' in globals() and topics:
        print(f"\n🎯 DISCOVERED TOPICS")
        for i, topic_words in enumerate(topics):
            print(f"Topic {i+1}: {', '.join(topic_words[:5])}")
    
    print("\n" + "="*60)
    print("END OF REPORT")
    print("="*60)

generate_comprehensive_report()

# ## 15. Export Results

def export_results_to_excel(filename="word_document_eda_results.xlsx"):
    """Export all analysis results to Excel file"""
    
    with pd.ExcelWriter(filename, engine='openpyxl') as writer:
        
        # Basic statistics
        analyzer.doc_stats.to_excel(writer, sheet_name='Document_Statistics')
        
        # Sentiment analysis
        if 'sentiment_results' in globals():
            sentiment_results.to_excel(writer, sheet_name='Sentiment_Analysis')
        
        # Readability analysis
        if 'readability_results' in globals():
            readability_results.to_excel(writer, sheet_name='Readability_Analysis')
        
        # Word frequencies
        freq_df = pd.DataFrame(dict([(doc, freq) for doc, freq in word_frequencies.items()])).fillna(0)
        freq_df.to_excel(writer, sheet_name='Word_Frequencies')
        
        # Document similarity
        if 'similarity_matrix' in globals() and similarity_matrix is not None:
            similarity_matrix.to_excel(writer, sheet_name='Document_Similarity')
    
    print(f"Results exported to {filename}")

# Export results
export_results_to_excel()

# ## 16. Usage Instructions for Your Own Documents

print("""
📝 TO USE WITH YOUR OWN DOCUMENTS:

1. SINGLE DOCUMENT:
   analyzer = WordDocumentAnalyzer()
   text = analyzer.load_single_document('path/to/your/document.docx')

2. MULTIPLE DOCUMENTS:
   file_paths = ['doc1.docx', 'doc2.docx', 'doc3.docx']
   documents = analyzer.load_multiple_documents(file_paths)

3. DIRECTORY OF DOCUMENTS:
   documents = analyzer.load_from_directory('path/to/directory', '*.docx')

4. Then run the analysis functions:
   - calculate_document_statistics()
   - perform_sentiment_analysis()
   - analyze_readability()
   - perform_topic_modeling()
   - analyze_document_similarity()

5. MODIFY PARAMETERS:
   - Change n_topics in topic modeling
   - Adjust top_n in word frequency analysis
   - Customize preprocessing steps
   - Add your own analysis functions

6. ADDITIONAL FEATURES YOU CAN ADD:
   - Named Entity Recognition (NER)
   - N-gram analysis
   - Keyword extraction
   - Language detection
   - Document classification
   - Time series analysis (if documents have dates)
""")

print("\n✅ Word Document EDA Analysis Complete!")
print("📊 All visualizations and analysis results are displayed above.")
print("📁 Results have been exported to Excel file.")
print("🔄 You can now modify the code to analyze your own documents.")
